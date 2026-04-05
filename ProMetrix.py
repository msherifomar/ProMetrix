"""
ProMetrix Analytics v2.2
Courtesy of Dr. M S Omar  BDS  MSc
For Doers, Researchers and Innovators
===========================================================================
Statistical Analysis & Publication-Quality Visualization

HOW TO RUN:
    %run ProMetrix.py          (in Spyder)
    python ProMetrix.py        (standalone)

REQUIRES:
    pip install pandas numpy scipy matplotlib seaborn openpyxl statsmodels PyQt5
    pip install python-docx          (optional, for Word export)
"""

import sys, os, warnings, tempfile
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Qt5Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.lines import Line2D
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg, NavigationToolbar2QT
import seaborn as sns
from scipy import stats
from scipy.stats import (pearsonr, shapiro, levene, kruskal, mannwhitneyu,
                          f_oneway, norm, f as f_dist)
from itertools import combinations

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QGroupBox, QLabel, QPushButton, QComboBox, QTextEdit, QTabWidget,
    QFileDialog, QMessageBox, QSplitter, QDialog, QDialogButtonBox,
    QFrame, QColorDialog, QGridLayout, QScrollArea, QProgressDialog
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QTextCursor, QColor, QIcon, QPixmap

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

warnings.filterwarnings('ignore')

# ── FONTS ──
plt.rcParams.update({
    'font.family': 'serif',
    'font.serif': ['DejaVu Serif', 'Times New Roman', 'Georgia', 'serif'],
    'font.size': 11, 'axes.titlesize': 16, 'axes.labelsize': 13,
    'figure.dpi': 150, 'axes.spines.top': False, 'axes.spines.right': False,
})
FT = {'family':'serif','size':18,'weight':'bold','color':'#111'}
FS = {'family':'serif','size':11,'style':'italic','color':'#555'}
FA = {'family':'serif','size':13}
FK = {'family':'serif','size':11}

COLOR_PRESETS = {
    'Default': ['#E57373','#FFB74D','#64B5F6','#81C784','#BA68C8',
                '#4DD0E1','#FF8A65','#A1887F','#90A4AE','#F06292'],
    'Clinical': ['#D32F2F','#1976D2','#388E3C','#F57C00','#7B1FA2',
                 '#0097A7','#C2185B','#455A64','#AFB42B','#5D4037'],
    'Pastel': ['#EF9A9A','#90CAF9','#A5D6A7','#FFCC80','#CE93D8',
               '#80DEEA','#FFAB91','#BCAAA4','#B0BEC5','#F48FB1'],
    'Grayscale': ['#212121','#616161','#9E9E9E','#BDBDBD','#757575',
                  '#424242','#E0E0E0','#505050','#888888','#333333'],
    'Journal': ['#0072B2','#D55E00','#009E73','#CC79A7','#F0E442',
                '#56B4E9','#E69F00','#000000','#999999','#882255'],
}
MARKERS = ['o','D','s','^','p','h','v','<','>','8']
SKIP = {'crown','sample','specimen','id','n','#','number','no','no.'}

STYLE_SHEET = """
QMainWindow { background: #F5F5F5; }
QGroupBox { font-weight: bold; font-size: 11px; border: 1px solid #CCC;
    border-radius: 6px; margin-top: 10px; padding-top: 14px; }
QGroupBox::title { subcontrol-origin: margin; left: 12px; padding: 0 6px; }
QPushButton#action { background: #222; color: white; border: none;
    border-radius: 5px; padding: 8px 16px; font-weight: bold; font-size: 11px; }
QPushButton#action:hover { background: #444; }
QPushButton#plot { background: #FFF; color: #333; border: 1px solid #BBB;
    border-radius: 5px; padding: 7px 14px; font-size: 11px; }
QPushButton#plot:hover { background: #EEE; border-color: #333; }
QPushButton#demo { background: #FFF3E0; color: #E65100; border: 1px solid #FFB74D;
    border-radius: 5px; padding: 7px 14px; font-weight: bold; font-size: 11px; }
QPushButton#demo:hover { background: #FFE0B2; }
QPushButton#table_btn { background: #555; color: white; border: none;
    border-radius: 5px; padding: 7px 14px; font-weight: bold; font-size: 11px; }
QPushButton#table_btn:hover { background: #777; }
QPushButton#manuscript { background: #1B5E20; color: white; border: none;
    border-radius: 5px; padding: 8px 16px; font-weight: bold; font-size: 11px; }
QPushButton#manuscript:hover { background: #2E7D32; }
QPushButton#export_full { background: #B71C1C; color: white; border: none;
    border-radius: 5px; padding: 10px 16px; font-weight: bold; font-size: 12px; }
QPushButton#export_full:hover { background: #D32F2F; }
QComboBox { padding: 5px 10px; border: 1px solid #BBB; border-radius: 4px;
    background: white; font-size: 11px; }
QTextEdit { font-family: Consolas, Courier New, monospace; font-size: 10px;
    background: #FAFAFA; border: 1px solid #DDD; border-radius: 4px; }
QTabWidget::pane { border: 1px solid #CCC; border-radius: 4px; }
QLabel#title { font-size: 22px; font-weight: bold; color: #222; }
QLabel#subtitle { font-size: 10px; color: #888; font-style: italic; }
QLabel#credit { font-size: 8px; color: #AAA; }
QLabel#status { font-size: 10px; }
"""


# ══════════ STATISTICAL FUNCTIONS ══════════
def welch_anova(*groups):
    k=len(groups); ns=np.array([len(g) for g in groups],float)
    ms=np.array([np.mean(g) for g in groups]); vs=np.array([np.var(g,ddof=1) for g in groups])
    vs[vs==0]=1e-10; w=ns/vs; sw=np.sum(w); gm=np.sum(w*ms)/sw
    num=np.sum(w*(ms-gm)**2)/(k-1); lam=np.sum((1-w/sw)**2/(ns-1))
    den=1+2*(k-2)/(k**2-1)*lam; F=num/den; df2=(k**2-1)/(3*lam)
    return F,k-1,df2,1-f_dist.cdf(F,k-1,df2)

def hedges_g(a, b):
    a,b=np.asarray(a,float),np.asarray(b,float); n1,n2=len(a),len(b)
    sp=np.sqrt(((n1-1)*np.var(a,ddof=1)+(n2-1)*np.var(b,ddof=1))/(n1+n2-2))
    if sp<1e-10: return 0.0
    return (np.mean(a)-np.mean(b))/sp*(1-3/(4*(n1+n2-2)-1))

def bootstrap_ci_g(a, b, n_boot=2000, seed=42):
    rng=np.random.RandomState(seed); gs=[]
    for _ in range(n_boot): gs.append(hedges_g(rng.choice(a,len(a),True),rng.choice(b,len(b),True)))
    return np.percentile(gs,[2.5,97.5])

def mde_two(n1,n2,sd_p,alpha=0.05,power=0.80):
    return (norm.ppf(1-alpha/2)+norm.ppf(power))*sd_p*np.sqrt(1/n1+1/n2)

def epsilon_sq(H,k,N): return (H-k+1)/(N-k)
def omega_sq(F,k,N): return ((k-1)*(F-1))/((k-1)*(F-1)+N)

def _sg(p):
    if p<.001: return '***'
    if p<.01: return '**'
    if p<.05: return '*'
    return 'ns'

def _wr(s, mx=14):
    if len(s)<=mx: return s
    w=s.replace('_',' ').split()
    if len(w)==1: return s
    m=len(w)//2; return ' '.join(w[:m])+'\n'+' '.join(w[m:])


# ══════════ FIGURE GENERATORS ══════════
def gen_raincloud(gd, names, colors, oc):
    k=len(names); arrs=[gd[g] for g in names]; H,pk=kruskal(*arrs); N=sum(len(a) for a in arrs)
    fig,ax=plt.subplots(figsize=(max(10,k*2.2),7),facecolor='white'); ax.set_facecolor('white')
    for idx,g in enumerate(names):
        v=gd[g]; c=colors[g]
        parts=ax.violinplot(v,positions=[idx],showmeans=False,showmedians=False,showextrema=False)
        for pc in parts['bodies']:
            m=np.mean(pc.get_paths()[0].vertices[:,0])
            pc.get_paths()[0].vertices[:,0]=np.clip(pc.get_paths()[0].vertices[:,0],m,np.inf)
            pc.set_facecolor(c); pc.set_alpha(.15); pc.set_edgecolor(c); pc.set_linewidth(.5)
        ax.boxplot([v],positions=[idx-.05],widths=.15,patch_artist=True,showfliers=False,zorder=3,
            boxprops=dict(facecolor=c,alpha=.7,lw=1.2,edgecolor='#333'),
            whiskerprops=dict(color='#333',lw=1.2),capprops=dict(color='#333',lw=1.2),
            medianprops=dict(color='white',lw=2))
        jit=np.random.uniform(-.18,-.08,len(v))
        ax.scatter(idx+jit,v,color=c,alpha=.6,s=40,zorder=4,edgecolors='white',linewidths=.5)
        ax.scatter(idx-.05,np.mean(v),marker='D',color='white',edgecolors='#333',s=60,zorder=5,linewidths=1.5)
    pairs=list(combinations(range(k),2)); nc=len(pairs); brk=[]
    for i,j in pairs:
        U,p=mannwhitneyu(gd[names[i]],gd[names[j]],alternative='two-sided'); pa=min(p*nc,1.)
        if pa<.001: brk.append((i,j,'***'))
        elif pa<.01: brk.append((i,j,'**'))
        elif pa<.05: brk.append((i,j,'*'))
    for i in range(k-1):
        U,p=mannwhitneyu(gd[names[i]],gd[names[i+1]],alternative='two-sided'); pa=min(p*nc,1.)
        if pa>=.05 and not any(b[:2]==(i,i+1) for b in brk): brk.append((i,i+1,'ns'))
    brk.sort(key=lambda x:(x[1]-x[0],x[0]))
    av=np.concatenate(arrs); yr=av.max()-av.min(); bs=av.max()+yr*.08; st=yr*.07
    for bi,(i,j,lb) in enumerate(brk):
        y=bs+bi*st; ax.plot([i,i,j,j],[y-st*.15,y,y,y-st*.15],color='#555',lw=1,zorder=6)
        ax.text((i+j)/2,y+st*.05,lb,ha='center',va='bottom',fontsize=10,family='serif',
            fontweight='bold' if lb!='ns' else 'normal',fontstyle='italic' if lb=='ns' else 'normal',
            color='#222' if lb!='ns' else '#999',zorder=6)
    ax.set_xlim(-.6,k-.3); ax.set_ylim(None,bs+len(brk)*st+yr*.05)
    ax.set_xticks(range(k)); ax.set_xticklabels([_wr(g) for g in names],**FK,fontweight='bold')
    ax.set_ylabel(oc,**FA); ax.set_title(oc,**FT,pad=25)
    ps='p < 0.001' if pk<.001 else 'p = {:.3f}'.format(pk)
    ax.text(.5,1.015,'Kruskal-Wallis: H({}, N={}) = {:.2f}, {}'.format(k-1,N,H,ps),transform=ax.transAxes,ha='center',**FS)
    leg=[mpatches.Patch(facecolor=colors[g],alpha=.7,edgecolor='#333',label=g) for g in names]
    leg.append(Line2D([0],[0],marker='D',color='w',markerfacecolor='white',markeredgecolor='#333',markersize=8,label='Mean'))
    ax.legend(handles=leg,loc='lower center',ncol=min(k+1,6),bbox_to_anchor=(.5,-.14),frameon=True,edgecolor='#CCC',prop={'family':'serif','size':9})
    ax.text(.5,-.21,'* p < 0.05    ** p < 0.01    *** p < 0.001    ns = not significant  (Bonferroni)',
        transform=ax.transAxes,ha='center',fontsize=8,family='serif',fontstyle='italic',color='#555')
    ax.grid(axis='y',alpha=.3,ls='--'); plt.tight_layout(); return fig

def gen_barplot(gd, names, colors, oc):
    fig,ax=plt.subplots(figsize=(max(8,len(names)*1.8),6),facecolor='white'); ax.set_facecolor('white')
    means=[np.mean(gd[g]) for g in names]; sds=[np.std(gd[g],ddof=1) for g in names]; cols=[colors[g] for g in names]
    ax.bar(range(len(names)),means,yerr=sds,capsize=5,color=cols,edgecolor='#333',lw=1,alpha=.8,zorder=3,
        error_kw=dict(lw=1.5,capthick=1.5,color='#333'))
    for i,g in enumerate(names):
        jit=np.random.uniform(-.15,.15,len(gd[g]))
        ax.scatter(i+jit,gd[g],color='#333',alpha=.5,s=25,zorder=4,edgecolors='white',linewidths=.5)
    for i,(m,s) in enumerate(zip(means,sds)):
        ax.text(i,m+s+max(means)*.02,'{:.1f}'.format(m),ha='center',va='bottom',fontsize=9,fontweight='bold',family='serif')
    ax.set_xticks(range(len(names))); ax.set_xticklabels([_wr(g) for g in names],**FK,fontweight='bold')
    ax.set_ylabel(oc,**FA); ax.set_title(oc,**FT,pad=15)
    ax.grid(axis='y',alpha=.3,ls='--'); plt.tight_layout(); return fig

def gen_scatter(groups, group_names, colors, xk, yk):
    fig,ax=plt.subplots(figsize=(10,8),facecolor='white'); ax.set_facecolor('white')
    ax_x,ax_y,vg=[],[],[]
    for i,g in enumerate(group_names):
        xv=groups[g].get(xk,[]); yv=groups[g].get(yk,[])
        if not xv or not yv: continue
        ml=min(len(xv),len(yv)); xv=np.array(xv[:ml]); yv=np.array(yv[:ml])
        c=colors[g]; mk=MARKERS[i%len(MARKERS)]
        ax.scatter(xv,yv,color=c,marker=mk,s=100,alpha=.7,edgecolors='white',linewidths=.8,zorder=4,label=g)
        if len(xv)>2: sl,ic=np.polyfit(xv,yv,1); xl=np.linspace(xv.min(),xv.max(),50); ax.plot(xl,sl*xl+ic,color=c,ls='--',alpha=.5,lw=1.5,zorder=3)
        ax_x.extend(xv); ax_y.extend(yv); vg.append(g)
    if len(ax_x)<3: plt.close(fig); return None
    ax_x=np.array(ax_x); ax_y=np.array(ax_y)
    sl,ic,ro,po,_=stats.linregress(ax_x,ax_y)
    xr=np.linspace(ax_x.min()-np.ptp(ax_x)*.05,ax_x.max()+np.ptp(ax_x)*.05,100); yp=sl*xr+ic
    ax.plot(xr,yp,color='#333',lw=2.5,zorder=5)
    n=len(ax_x); xm=ax_x.mean(); se=np.sqrt(((ax_y-(sl*ax_x+ic))**2).sum()/(n-2))
    ci=1.96*se*np.sqrt(1/n+(xr-xm)**2/((ax_x-xm)**2).sum())
    ax.fill_between(xr,yp-ci,yp+ci,alpha=.15,color='gray',zorder=2)
    ps='P < .001' if po<.001 else 'P = {:.3f}'.format(po)
    ax.text(.03,.97,'Overall: r = {:.3f},  R\u00b2 = {:.3f},  {}'.format(ro,ro**2,ps),
        transform=ax.transAxes,fontsize=12,fontweight='bold',va='top',family='serif',
        bbox=dict(boxstyle='round,pad=0.5',fc='white',ec='#333',alpha=.9))
    wt=[]
    for g in vg:
        xv=groups[g].get(xk,[]); yv=groups[g].get(yk,[]); ml=min(len(xv),len(yv))
        if ml>2:
            r,p=pearsonr(xv[:ml],yv[:ml]); s='(***)' if p<.001 else '(**)' if p<.01 else '(*)' if p<.05 else '(ns)'
            wt.append('{}: r = {:.2f} {}'.format(g,r,s))
    if wt: ax.text(.03,.82,'\n'.join(wt),transform=ax.transAxes,fontsize=9,va='top',family='serif',
        bbox=dict(boxstyle='round,pad=0.4',fc='white',ec='#AAA',alpha=.85))
    ax.set_xlabel(xk,**FA,fontweight='bold'); ax.set_ylabel(yk,**FA,fontweight='bold')
    ax.set_title('{} vs {}'.format(xk,yk),**FT,pad=15)
    ax.legend(loc='lower right',frameon=True,edgecolor='#CCC',prop={'family':'serif','size':10})
    ax.grid(alpha=.3,ls='--'); plt.tight_layout(); return fig

def gen_heatmap(groups, group_names, outcomes):
    good=[oc for oc in outcomes if sum(len(groups[g].get(oc,[])) for g in group_names)>=5]
    if len(good)<2: return None
    rd=[]
    for g in group_names:
        mx=max((len(groups[g].get(oc,[])) for oc in good),default=0)
        for i in range(mx):
            row={}; ok=True
            for oc in good:
                vs=groups[g].get(oc,[])
                if i<len(vs): row[oc]=vs[i]
                else: ok=False; break
            if ok: rd.append(row)
    df=pd.DataFrame(rd)
    if len(df)<5: return None
    nv=len(good); fig,ax=plt.subplots(figsize=(max(8,nv*1.5),max(6,nv*1.2)),facecolor='white')
    corr=df[good].corr(); annot=np.empty_like(corr,dtype=object)
    for i in range(nv):
        for j in range(nv):
            if i==j: annot[i,j]=''
            else:
                r,p=pearsonr(df[good[i]],df[good[j]]); s='***' if p<.001 else '**' if p<.01 else '*' if p<.05 else ''
                annot[i,j]='{:.2f}{}'.format(r,s)
    mask=np.triu(np.ones_like(corr,dtype=bool),k=1)
    sns.heatmap(corr,mask=mask,annot=annot,fmt='',cmap='RdBu_r',center=0,vmin=-1,vmax=1,
        square=True,linewidths=1,linecolor='white',ax=ax,cbar_kws={'shrink':.8,'label':'Pearson r'},
        annot_kws={'size':11,'fontweight':'bold','family':'serif'})
    ax.set_title('Correlation Matrix (Pooled)',**FT,pad=15)
    ax.set_xticklabels(ax.get_xticklabels(),rotation=45,ha='right',fontsize=9,family='serif')
    ax.set_yticklabels(ax.get_yticklabels(),rotation=0,fontsize=9,family='serif')
    plt.tight_layout(); return fig

def gen_table_fig(cl, rl, cd, title, hl=None, sec=None):
    nr=len(rl); nc_=len(cl); fw=max(8,nc_*1.7+2.5); fh=max(2.5,(nr+1)*.32+1.5)
    fig,ax=plt.subplots(figsize=(fw,fh),facecolor='white')
    ax.set_facecolor('white'); ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1)
    ax.text(.02,.97,title,fontsize=11,fontweight='bold',va='top',ha='left',family='serif',style='italic')
    rh=min(.055,.82/(nr+1)); ty=.90; lx=.02; tw=.96
    lw=min(.22,max(.12,max(len(str(r)) for r in rl)*.011)); dw=(tw-lw)/nc_
    ax.plot([lx,lx+tw],[ty,ty],color='black',lw=1.8,clip_on=False)
    x=lx+lw
    for c in cl: ax.text(x+dw*.5,ty-rh*.5,c,ha='center',va='center',fontsize=9.5,fontweight='bold',family='serif'); x+=dw
    bhy=ty-rh; ax.plot([lx,lx+tw],[bhy,bhy],color='black',lw=.8,clip_on=False)
    for ri in range(nr):
        y=bhy-(ri+.5)*rh
        if sec and ri in sec:
            ax.text(lx+.01,y,rl[ri],ha='left',va='center',fontsize=9.5,fontweight='bold',family='serif',style='italic')
            if ri>0: ax.plot([lx,lx+tw],[bhy-ri*rh,bhy-ri*rh],color='black',lw=.4,clip_on=False)
            continue
        ax.text(lx+.01,y,str(rl[ri]),ha='left',va='center',fontsize=9,family='serif')
        x=lx+lw
        for ci in range(nc_):
            txt=str(cd[ri][ci]) if ri<len(cd) and ci<len(cd[ri]) else ''
            props=dict(ha='center',va='center',fontsize=9,family='serif',color='black')
            if hl and (ri,ci) in hl: props['fontweight']='bold'
            ax.text(x+dw*.5,y,txt,**props); x+=dw
    ax.plot([lx,lx+tw],[bhy-nr*rh,bhy-nr*rh],color='black',lw=1.8,clip_on=False)
    plt.subplots_adjust(left=0,right=1,top=1,bottom=0); return fig


# ══════════ ANALYSIS ENGINE ══════════
def analyze_outcome(gd, names, oc):
    arrs=[gd[g] for g in names]; k=len(names); N=sum(len(a) for a in arrs)
    norm_res={}; all_normal=True
    for g in names:
        if len(gd[g])>=3: _,p=shapiro(gd[g]); norm_res[g]=p
        if norm_res.get(g,1)<=0.05: all_normal=False
    L,pl=levene(*arrs); eq=pl>0.05; nonp=not all_normal or not eq
    F_val,pa=f_oneway(*arrs); H_val,pk=kruskal(*arrs)
    Fw,d1w,d2w,pw=welch_anova(*arrs)
    om=omega_sq(F_val,k,N); es=epsilon_sq(H_val,k,N); agree=(pk<0.05)==(pw<0.05)
    pairs=list(combinations(range(k),2)); ncp=len(pairs); pw_res=[]
    for i,j in pairs:
        g1,g2=names[i],names[j]; a1,a2=gd[g1],gd[g2]; md=np.mean(a1)-np.mean(a2)
        U,pr=mannwhitneyu(a1,a2,alternative='two-sided'); padj=min(pr*ncp,1.)
        n1,n2=len(a1),len(a2); z=(U-n1*n2/2)/np.sqrt(n1*n2*(n1+n2+1)/12)
        rb=abs(z)/np.sqrt(n1+n2); gv=hedges_g(a1,a2); ci=bootstrap_ci_g(a1,a2)
        sp=np.sqrt((np.var(a1,ddof=1)+np.var(a2,ddof=1))/2); mv=mde_two(n1,n2,sp)
        pw_res.append(dict(g1=g1,g2=g2,md=md,padj=padj,rb=rb,hg=gv,ci_lo=ci[0],ci_hi=ci[1],mde=mv))
    non_normal=[g for g in names if norm_res.get(g,1)<=0.05]
    return dict(oc=oc,names=names,gd=gd,N=N,k=k,nonp=nonp,all_normal=all_normal,eq_var=eq,p_levene=pl,
        non_normal_groups=non_normal,H=H_val,pk=pk,es=es,F=F_val,pa=pa,om=om,
        Fw=Fw,df1w=d1w,df2w=d2w,pw_welch=pw,agree=agree,pairs=pw_res,ncp=ncp)

def gen_methods_text(analyses):
    a0=analyses[0]; k=a0['k']; oc_names=[a['oc'] for a in analyses]; ncp=a0['ncp']
    L=[]; L.append("Statistical Analysis"); L.append("")
    L.append("Statistical analyses were performed using ProMetrix Analytics (v2.2). "
        "Descriptive statistics including means, standard deviations, medians, interquartile ranges, "
        "and coefficients of variation (CV%) were calculated for each group and outcome variable. "
        "The outcome variables analyzed were: {}.".format('; '.join(oc_names))); L.append("")
    L.append("Normality of data distribution was assessed for each group using the Shapiro-Wilk test, "
        "and homogeneity of variances was evaluated using Levene's test. When violations of parametric "
        "assumptions were detected (non-normal distributions or unequal variances), the Kruskal-Wallis "
        "test was employed as the primary omnibus test. When parametric assumptions were satisfied, "
        "one-way ANOVA was used."); L.append("")
    L.append("Post-hoc pairwise comparisons were conducted using Mann-Whitney U tests (for nonparametric "
        "analyses) or Welch's t-tests (for parametric analyses), with Bonferroni correction for {} "
        "comparisons per outcome (adjusted alpha = {:.4f}). Effect sizes were reported as Hedges' g "
        "with bias correction and 95% bootstrap confidence intervals (2,000 resamples). The rank-biserial "
        "correlation (r) was additionally reported as a nonparametric effect size measure.".format(ncp,0.05/ncp)); L.append("")
    L.append("Welch's ANOVA was conducted as a sensitivity analysis for each outcome to verify robustness "
        "of findings across parametric and nonparametric approaches. The minimum detectable effect (MDE) "
        "at 80% power was calculated for each pairwise comparison to contextualize the study's statistical "
        "sensitivity. Pearson correlation coefficients were computed to examine relationships between "
        "outcome variables. Statistical significance was set at alpha = 0.05 for all tests.")
    return '\n'.join(L)

def gen_results_text(analyses):
    L=["Results",""]
    for a in analyses:
        oc=a['oc']; names=a['names']; k=a['k']; N=a['N']
        desc=["{} ({:.2f} +/- {:.2f})".format(g,np.mean(a['gd'][g]),np.std(a['gd'][g],ddof=1)) for g in names]
        L.append("{}: Mean values (mean +/- SD) were: {}.".format(oc,'; '.join(desc))); L.append("")
        if a['nonp']:
            ps='p < 0.001' if a['pk']<0.001 else 'p = {:.3f}'.format(a['pk'])
            L.append("{} differed significantly across groups (Kruskal-Wallis H({}) = {:.2f}, {}, epsilon-squared = {:.3f}).".format(oc,k-1,a['H'],ps,a['es']))
        else:
            ps='p < 0.001' if a['pa']<0.001 else 'p = {:.3f}'.format(a['pa'])
            L.append("{} differed significantly across groups (F({}, {}) = {:.2f}, {}, omega-squared = {:.3f}).".format(oc,k-1,N-k,a['F'],ps,a['om']))
        L.append("")
        sig=[r for r in a['pairs'] if r['padj']<0.05]; ns_p=[r for r in a['pairs'] if r['padj']>=0.05]
        if sig:
            parts=["{} vs {} (Hedges' g = {:+.2f}, 95% CI [{:+.2f}, {:+.2f}], {})".format(
                r['g1'],r['g2'],r['hg'],r['ci_lo'],r['ci_hi'],'p < 0.001' if r['padj']<0.001 else 'p = {:.3f}'.format(r['padj'])) for r in sig]
            L.append("Post-hoc comparisons revealed significant differences between: "+'; '.join(parts)+".")
        if ns_p: L.append("No significant differences were observed between: "+'; '.join(['{} vs {}'.format(r['g1'],r['g2']) for r in ns_p])+" (all adjusted p > 0.05).")
        L.append("")
        pws='p < 0.001' if a['pw_welch']<0.001 else 'p = {:.3f}'.format(a['pw_welch'])
        L.append("Welch's ANOVA {} the primary analysis (F({:.0f}, {:.1f}) = {:.2f}, {}). Results were {} across parametric and nonparametric approaches.".format(
            'confirmed' if a['agree'] else 'showed discrepant results compared to',a['df1w'],a['df2w'],a['Fw'],pws,'consistent' if a['agree'] else 'inconsistent'))
        L.append("")
        avg_mde=np.mean([r['mde'] for r in a['pairs']])
        L.append("With n = {} per group, the average minimum detectable pairwise difference at 80% power for {} was {:.2f}.".format(int(N/k),oc,avg_mde))
        if sig:
            best=max(sig,key=lambda r:abs(r['hg']))
            L.append("The largest observed effect was between {} and {} (Hedges' g = {:+.2f}).".format(best['g1'],best['g2'],best['hg']))
        L.append("")
    return '\n'.join(L)


# ══════════ GUI ══════════
class ColorButton(QPushButton):
    def __init__(self,gn,color,cb):
        super().__init__(); self.gn=gn; self.color=color; self.cb=cb
        self.setFixedSize(24,24); self._us(); self.clicked.connect(self._pk)
    def _us(self):
        self.setStyleSheet("QPushButton{{background:{};border:2px solid #555;border-radius:4px;}}"
            "QPushButton:hover{{border:2px solid #000;}}".format(self.color))
    def _pk(self):
        c=QColorDialog.getColor(QColor(self.color),self.parent())
        if c.isValid(): self.color=c.name(); self._us(); self.cb(self.gn,self.color)


class ProMetrix(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("ProMetrix Analytics")
        self.setMinimumSize(1150,700); self.resize(1400,920)
        self.setStyleSheet(STYLE_SHEET)
        # Set icon
        icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),'prometrix_icon.ico')
        if os.path.exists(icon_path): self.setWindowIcon(QIcon(icon_path))
        self.groups={}; self.group_names=[]; self.outcomes=[]
        self.colors={}; self.fig=None; self.color_buttons={}
        self._last_table_data=None; self._last_analysis=None
        self._build(); print("[ProMetrix Analytics] Ready.")

    def _build(self):
        central=QWidget(); self.setCentralWidget(central)
        ml=QHBoxLayout(central); ml.setContentsMargins(8,8,8,8)
        sp=QSplitter(Qt.Horizontal); ml.addWidget(sp)
        scroll=QScrollArea(); scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff); scroll.setFixedWidth(390)
        left=QWidget(); lv=QVBoxLayout(left); lv.setContentsMargins(5,5,5,5)
        t=QLabel("ProMetrix Analytics"); t.setObjectName("title"); t.setAlignment(Qt.AlignCenter); lv.addWidget(t)
        s=QLabel("For Doers, Researchers and Innovators")
        s.setObjectName("subtitle"); s.setAlignment(Qt.AlignCenter); lv.addWidget(s)
        cr=QLabel("Courtesy of Dr. M S Omar  BDS  MSc")
        cr.setObjectName("credit"); cr.setAlignment(Qt.AlignCenter); lv.addWidget(cr)
        lv.addWidget(self._sep())

        g1=QGroupBox(" 1. Load Data"); g1v=QVBoxLayout(g1)
        g1v.addWidget(QLabel("Excel (each sheet = group)\nor CSV (with 'Group' column)"))
        bh=QHBoxLayout()
        bx=QPushButton("Load Excel"); bx.setObjectName("action"); bx.clicked.connect(self._load_excel)
        bc=QPushButton("Load CSV"); bc.setObjectName("action"); bc.clicked.connect(self._load_csv)
        bh.addWidget(bx); bh.addWidget(bc); g1v.addLayout(bh)
        bd=QPushButton("Load Demo Data (Crown Study)"); bd.setObjectName("demo"); bd.clicked.connect(self._load_demo); g1v.addWidget(bd)
        self.lbl=QLabel("No data loaded"); self.lbl.setObjectName("status"); self.lbl.setStyleSheet("color:#999;"); g1v.addWidget(self.lbl); lv.addWidget(g1)

        g2=QGroupBox(" 2. Select Outcome"); g2v=QVBoxLayout(g2)
        self.oc_box=QComboBox(); g2v.addWidget(self.oc_box); lv.addWidget(g2)

        gc=QGroupBox(" 3. Colors"); gcl=QVBoxLayout(gc)
        ph=QHBoxLayout(); ph.addWidget(QLabel("Preset:"))
        self.preset_box=QComboBox(); self.preset_box.addItems(list(COLOR_PRESETS.keys()))
        self.preset_box.currentTextChanged.connect(self._apply_preset); ph.addWidget(self.preset_box)
        gcl.addLayout(ph); self.cgrid=QGridLayout(); gcl.addLayout(self.cgrid); lv.addWidget(gc)

        g4=QGroupBox(" 4. Analyze & Plot"); g4v=QVBoxLayout(g4)
        bs=QPushButton("Run Full Statistics"); bs.setObjectName("action"); bs.clicked.connect(self._run_stats); g4v.addWidget(bs)
        g4v.addWidget(self._sep()); g4v.addWidget(QLabel("Figures:"))
        for row in [[("Raincloud",self._raincloud),("Bar Plot",self._barplot)],[("Scatter",self._scatter_dlg),("Heatmap",self._heatmap)]]:
            rh=QHBoxLayout()
            for txt,fn in row: b=QPushButton(txt); b.setObjectName("plot"); b.clicked.connect(fn); rh.addWidget(b)
            g4v.addLayout(rh)
        g4v.addWidget(self._sep()); g4v.addWidget(QLabel("Tables:"))
        for row in [[("Descriptive",self._tbl_desc),("Pairwise",self._tbl_pw)],[("Correlation",self._tbl_corr),("CV%",self._tbl_cv)]]:
            rh=QHBoxLayout()
            for txt,fn in row: b=QPushButton(txt); b.setObjectName("table_btn"); b.clicked.connect(fn); rh.addWidget(b)
            g4v.addLayout(rh)
        g4v.addWidget(self._sep())
        bm=QPushButton("Generate Manuscript Text"); bm.setObjectName("manuscript"); bm.clicked.connect(self._manuscript); g4v.addWidget(bm)
        g4v.addWidget(self._sep())
        bsv=QPushButton("Save Current Figure (PNG / PDF)"); bsv.setObjectName("action"); bsv.clicked.connect(self._save_fig); g4v.addWidget(bsv)
        bfull=QPushButton("EXPORT FULL REPORT (.docx)"); bfull.setObjectName("export_full"); bfull.clicked.connect(self._export_full); g4v.addWidget(bfull)
        be=QPushButton("Export Stats to CSV"); be.setObjectName("action"); be.clicked.connect(self._export); g4v.addWidget(be)
        lv.addWidget(g4); lv.addStretch(); scroll.setWidget(left)

        right=QWidget(); rv=QVBoxLayout(right); rv.setContentsMargins(0,0,0,0)
        self.tabs=QTabWidget(); rv.addWidget(self.tabs)
        self.txt=QTextEdit(); self.txt.setReadOnly(True); self.tabs.addTab(self.txt,"Statistics")
        self.fig_widget=QWidget(); self.fig_layout=QVBoxLayout(self.fig_widget)
        self.fig_layout.setContentsMargins(0,0,0,0); self.tabs.addTab(self.fig_widget,"Figure / Table")
        self.ms_txt=QTextEdit(); self.ms_txt.setReadOnly(True)
        self.ms_txt.setStyleSheet("font-family:'Times New Roman',serif;font-size:12px;background:white;")
        self.tabs.addTab(self.ms_txt,"Manuscript Text")
        sp.addWidget(scroll); sp.addWidget(right); sp.setStretchFactor(0,0); sp.setStretchFactor(1,1)
        self._welcome()

    def _sep(self): f=QFrame(); f.setFrameShape(QFrame.HLine); return f

    # ════ COLORS ════
    def _rebuild_cs(self):
        while self.cgrid.count():
            w=self.cgrid.takeAt(0).widget()
            if w: w.deleteLater()
        self.color_buttons={}
        for i,g in enumerate(self.group_names):
            btn=ColorButton(g,self.colors[g],self._occ); lbl=QLabel(g); lbl.setStyleSheet("font-size:10px;")
            self.cgrid.addWidget(btn,i,0); self.cgrid.addWidget(lbl,i,1); self.color_buttons[g]=btn
    def _occ(self,gn,nc): self.colors[gn]=nc
    def _apply_preset(self,pn):
        if pn not in COLOR_PRESETS: return
        pal=COLOR_PRESETS[pn]
        for i,g in enumerate(self.group_names):
            self.colors[g]=pal[i%len(pal)]
            if g in self.color_buttons: self.color_buttons[g].color=self.colors[g]; self.color_buttons[g]._us()
    def _assign_colors(self):
        pal=COLOR_PRESETS.get(self.preset_box.currentText(),COLOR_PRESETS['Default'])
        self.colors={g:pal[i%len(pal)] for i,g in enumerate(self.group_names)}

    # ════ DATA ════
    def _load_excel(self):
        p,_=QFileDialog.getOpenFileName(self,"Select Excel","","Excel (*.xlsx *.xls);;All (*)")
        if not p: return
        try:
            xls=pd.ExcelFile(p); sheets={s:pd.read_excel(p,sheet_name=s) for s in xls.sheet_names}
            first=list(sheets.values())[0]; gc=self._fgc(first)
            if gc: self._fl(first,gc)
            else: self._fs(sheets)
            self._pl()
        except Exception as e: QMessageBox.critical(self,"Error",str(e))
    def _load_csv(self):
        p,_=QFileDialog.getOpenFileName(self,"Select CSV","","CSV (*.csv);;All (*)")
        if not p: return
        try: df=pd.read_csv(p); gc=self._fgc(df) or df.columns[0]; self._fl(df,gc); self._pl()
        except Exception as e: QMessageBox.critical(self,"Error",str(e))
    def _load_demo(self):
        self.groups={
            'Human Expert':{'Design Time (s)':[144,138,121,158,105,121,99,89,82,94],'Volume Removed (mm3)':[18.18,8.54,16.56,20.52,15.38,20.06,13.65,14.82,12.15,13.22],'Total Adj Time (s)':[410,217,310,341,338,385,83,245,305,293],'Surface Deviation (um)':[50.5,34.0,68.5,60.1,47.9,59.7,37.9,47.8,37.7,40.0]},
            'Human Novice':{'Design Time (s)':[155,165,208,88,182,112,85,108,88,82],'Volume Removed (mm3)':[7.05,13.90,32.87,56.52,9.57,6.39,4.77,8.87,17.65,14.35],'Total Adj Time (s)':[54,404,448,740,115,71,116,124,115,59],'Surface Deviation (um)':[]},
            '3Shape Automate':{'Design Time (s)':[91,90,92,91,94,88,90,87,89,90],'Volume Removed (mm3)':[23.89,45.12,34.75,29.26,51.31,23.73,14.72,5.72,38.70,32.35],'Total Adj Time (s)':[1491,1458,917,1243,1182,480,230,246,963,620],'Surface Deviation (um)':[96.6,175.4,122.6,101.5,193.7,102.2,64.1,42.9,162.7,123.6]},
            'ExoCAD AI':{'Design Time (s)':[91,100,92,91,96,90,91,91,91,90],'Volume Removed (mm3)':[11.76,17.51,4.76,11.24,13.49,9.28,9.38,10.02,9.43,10.51],'Total Adj Time (s)':[230,308,158,207,172,116,179,168,246,173],'Surface Deviation (um)':[63.7,73.9,41.2,56.3,54.0,39.0,52.1,41.2,46.4,41.0]},
            'Medit AI':{'Design Time (s)':[16,14,15,15,16,15,16,14,15,17],'Volume Removed (mm3)':[18.11,7.32,15.25,11.42,23.45,32.47,40.05,2.30,1.42,23.40],'Total Adj Time (s)':[225,65,215,169,526,441,492,65,63,323],'Surface Deviation (um)':[124.6,26.5,59.3,37.3,109.0,157.3,193.8,27.6,25.1,107.8]},
        }
        self.group_names=list(self.groups.keys())
        self.outcomes=['Design Time (s)','Volume Removed (mm3)','Total Adj Time (s)','Surface Deviation (um)']
        self._assign_colors(); self._pl()
    @staticmethod
    def _fgc(df):
        for c in df.columns:
            if c.lower() in ('group','groups','platform','method','type','category'): return c
        return None
    def _nc(self,df,ex=None):
        skip=SKIP|({ex.lower()} if ex else set())
        return [c for c in df.columns if c.lower() not in skip and pd.api.types.is_numeric_dtype(df[c])]
    def _fs(self,sheets):
        self.group_names=list(sheets.keys()); common=None
        for df in sheets.values(): nc=set(self._nc(df)); common=nc if common is None else common&nc
        self.outcomes=sorted(common,key=lambda x:list(sheets.values())[0].columns.tolist().index(x) if x in list(sheets.values())[0].columns else 999)
        self.groups={g:{c:df[c].dropna().astype(float).tolist() for c in self.outcomes} for g,df in sheets.items()}
        self._assign_colors()
    def _fl(self,df,gc):
        self.group_names=df[gc].unique().tolist(); self.outcomes=self._nc(df,gc)
        self.groups={g:{c:df[df[gc]==g][c].dropna().astype(float).tolist() for c in self.outcomes} for g in self.group_names}
        self._assign_colors()
    def _pl(self):
        self.lbl.setText("Loaded: {} groups, {} outcomes".format(len(self.group_names),len(self.outcomes)))
        self.lbl.setStyleSheet("color:#2E7D32;font-weight:bold;")
        self.oc_box.clear(); self.oc_box.addItems(self.outcomes); self._rebuild_cs(); self._summary()

    # ════ HELPERS ════
    def _ok(self):
        if not self.groups: QMessageBox.warning(self,"","Load data first."); return False
        return True
    def _gd(self,oc):
        gd={g:np.array(self.groups[g].get(oc,[]),dtype=float) for g in self.group_names if self.groups[g].get(oc,[])}
        if len(gd)<2: QMessageBox.warning(self,"","Need 2+ groups."); return None
        return gd
    def _w(self,t): self.txt.moveCursor(QTextCursor.End); self.txt.insertPlainText(t)
    def _show(self,fig):
        while self.fig_layout.count():
            w=self.fig_layout.takeAt(0).widget()
            if w: w.deleteLater()
        self.fig=fig; canvas=FigureCanvasQTAgg(fig); toolbar=NavigationToolbar2QT(canvas,self.fig_widget)
        self.fig_layout.addWidget(toolbar); self.fig_layout.addWidget(canvas); canvas.draw(); self.tabs.setCurrentIndex(1)

    # ════ RUN STATS ════
    def _run_stats(self):
        if not self._ok(): return
        oc=self.oc_box.currentText(); gd=self._gd(oc)
        if not gd: return
        a=analyze_outcome(gd,list(gd.keys()),oc); self._last_analysis=a
        names=a['names']; k=a['k']; N=a['N']
        self.txt.clear(); self.tabs.setCurrentIndex(0); W=self._w
        W("="*70+"\n   ProMetrix Analytics: {}\n   {} groups  |  N = {}\n".format(oc,k,N)+"="*70+"\n\n")
        W("DESCRIPTIVE STATISTICS\n{:<22} {:>4} {:>12} {:>12} {:>8}\n".format('Group','n','Mean +/- SD','Median','CV%'))
        W("-"*60+"\n")
        for g in names:
            arr=gd[g]; m=np.mean(arr); s=np.std(arr,ddof=1); cv=s/m*100 if m else 0
            W("{:<22} {:>4} {:>12} {:>12} {:>7.1f}%{}\n".format(g,len(arr),'{:.2f} +/- {:.2f}'.format(m,s),'{:.2f}'.format(np.median(arr)),cv,"  !" if cv>50 else ""))
        W("\nASSUMPTION TESTING\n")
        for g in names:
            if len(gd[g])>=3: _,p=shapiro(gd[g]); W("  Shapiro  {:<18} p={:.4f}  {}\n".format(g,p,'Normal' if p>.05 else 'NOT normal'))
        W("  Levene   {:<18} p={:.4f}  {}\n".format('(all)',a['p_levene'],'Equal' if a['eq_var'] else 'UNEQUAL'))
        W("\n"+"="*70+"\n  PRIMARY: {}\n  SENSITIVITY: Welch ANOVA\n".format(
            'Kruskal-Wallis + Mann-Whitney U (Bonferroni)' if a['nonp'] else 'ANOVA + Welch t (Bonferroni)')+"="*70+"\n\n")
        W("OMNIBUS\n  ANOVA  F({},{})={:.3f} p={:.6f} {} omega2={:.3f}\n".format(k-1,N-k,a['F'],a['pa'],_sg(a['pa']),a['om']))
        W("  KW    H({})={:.3f}    p={:.6f} {} eps2={:.3f}\n".format(k-1,a['H'],a['pk'],_sg(a['pk']),a['es']))
        W("  Welch F({},{:.1f})={:.3f} p={:.6f} {}\n".format(a['df1w'],a['df2w'],a['Fw'],a['pw_welch'],_sg(a['pw_welch'])))
        W("  >> {}\n".format('Results AGREE (robust)' if a['agree'] else 'Results DISAGREE'))
        W("\nPAIRWISE ({} pairs, Bonferroni)\n".format(a['ncp'])+"-"*95+"\n")
        W("{:<30} {:>9} {:>5}   {:>22}   {:>8}\n".format('Comparison','p-adj','Sig','Hedges g [95% CI]','MDE'))
        W("-"*95+"\n")
        for r in a['pairs']:
            W("{:<30} {:>9.4f} {:>5}   {:>+6.2f} [{:>+6.2f}, {:>+6.2f}]   {:>8.2f}\n".format(
                '{} vs {}'.format(r['g1'],r['g2']),r['padj'],_sg(r['padj']),r['hg'],r['ci_lo'],r['ci_hi'],r['mde']))
        W("-"*95+"\n")

    # ════ MANUSCRIPT ════
    def _manuscript(self):
        if not self._ok(): return
        analyses=[]
        for oc in self.outcomes:
            gd=self._gd(oc)
            if gd: analyses.append(analyze_outcome(gd,list(gd.keys()),oc))
        if not analyses: return
        self.ms_txt.clear()
        self.ms_txt.setPlainText(gen_methods_text(analyses)+'\n\n'+'-'*60+'\n\n'+gen_results_text(analyses))
        self.tabs.setCurrentIndex(2)

    # ════ EXPORT FULL REPORT ════
    def _export_full(self):
        if not self._ok(): return
        if not HAS_DOCX:
            return QMessageBox.warning(self,"Missing Package","python-docx not installed.\n\nRun:\n    pip install python-docx")
        p,_=QFileDialog.getSaveFileName(self,"Export Full Report","ProMetrix_Report.docx","Word (*.docx)")
        if not p: return
        prog=QProgressDialog("Generating full report...","",0,100,self)
        prog.setWindowModality(Qt.WindowModal); prog.setMinimumDuration(0); prog.setValue(0)
        QApplication.processEvents(); tmp_files=[]
        try:
            doc=Document()
            # BODY: TNR 12pt double-spaced
            style=doc.styles['Normal']; style.font.name='Times New Roman'; style.font.size=Pt(12)
            style.paragraph_format.line_spacing_rule=WD_LINE_SPACING.DOUBLE
            style.paragraph_format.space_after=Pt(0)
            for sec in doc.sections:
                sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54)
                sec.left_margin=Cm(2.54); sec.right_margin=Cm(2.54)

            def add_h(text,level=1):
                h=doc.add_heading(text,level=level)
                for r in h.runs: r.font.name='Times New Roman'; r.font.size=Pt(14 if level==1 else 12)
                h.paragraph_format.line_spacing_rule=WD_LINE_SPACING.DOUBLE; return h
            def add_p(text,bold=False,italic=False):
                pp=doc.add_paragraph(); pp.paragraph_format.line_spacing_rule=WD_LINE_SPACING.DOUBLE
                pp.paragraph_format.space_after=Pt(0)
                r=pp.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(12)
                if bold: r.bold=True
                if italic: r.italic=True
                return pp
            def add_fig(fig,caption=''):
                tmp=tempfile.NamedTemporaryFile(suffix='.png',delete=False)
                fig.savefig(tmp.name,dpi=300,bbox_inches='tight',facecolor='white',pad_inches=0.1)
                plt.close(fig); tmp_files.append(tmp.name)
                doc.add_picture(tmp.name,width=Inches(6))
                doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    cp.paragraph_format.line_spacing_rule=WD_LINE_SPACING.DOUBLE
                    r=cp.add_run(caption); r.font.name='Times New Roman'; r.font.size=Pt(10); r.italic=True
                add_p('')

            def add_tbl(cl,rl,cd,title,sec_rows=None):
                # TABLES: TNR 10pt, single-spaced
                tp=doc.add_paragraph()
                tp.paragraph_format.line_spacing_rule=WD_LINE_SPACING.SINGLE
                r=tp.add_run(title); r.bold=True; r.italic=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
                ncols=len(cl)+1; nrows=len(rl)
                table=doc.add_table(rows=1+nrows,cols=ncols)
                table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
                def _brd(cell,top=None,bot=None):
                    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
                    tcB=tcPr.find(qn('w:tcBorders'))
                    if tcB is None: tcB=tc.makeelement(qn('w:tcBorders'),{}); tcPr.append(tcB)
                    for edge,val in [('top',top),('bottom',bot),('left',None),('right',None)]:
                        el=tcB.find(qn('w:{}'.format(edge)))
                        if el is None: el=tc.makeelement(qn('w:{}'.format(edge)),{}); tcB.append(el)
                        if val: el.set(qn('w:val'),'single'); el.set(qn('w:sz'),val); el.set(qn('w:color'),'000000')
                        else: el.set(qn('w:val'),'none')
                hc=table.rows[0].cells; hc[0].text=''
                for ci,c in enumerate(cl): hc[ci+1].text=c
                for ci in range(ncols):
                    for pp in hc[ci].paragraphs:
                        pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                        pp.paragraph_format.line_spacing_rule=WD_LINE_SPACING.SINGLE
                        for r in pp.runs: r.bold=True; r.font.size=Pt(10); r.font.name='Times New Roman'
                    _brd(hc[ci],top='12',bot='6')
                for ri in range(nrows):
                    rc=table.rows[ri+1].cells; is_sec=sec_rows and ri in sec_rows; is_last=ri==nrows-1
                    rc[0].text=str(rl[ri])
                    if ri<len(cd):
                        for ci in range(min(len(cd[ri]),len(cl))): rc[ci+1].text=str(cd[ri][ci])
                    for ci in range(ncols):
                        for pp in rc[ci].paragraphs:
                            pp.alignment=WD_ALIGN_PARAGRAPH.CENTER if ci>0 else WD_ALIGN_PARAGRAPH.LEFT
                            pp.paragraph_format.line_spacing_rule=WD_LINE_SPACING.SINGLE
                            for r in pp.runs:
                                r.font.size=Pt(10); r.font.name='Times New Roman'
                                if is_sec: r.bold=True; r.italic=True
                        _brd(rc[ci],top='4' if (sec_rows and ri in sec_rows and ri>0) else None,bot='12' if is_last else None)
                add_p('')

            # ─── BUILD REPORT ───
            prog.setValue(5)
            add_h('ProMetrix Analytics Report',1)
            add_p('Courtesy of Dr. M S Omar  BDS  MSc',italic=True)
            add_p('')

            prog.setLabelText("Analyzing all outcomes..."); QApplication.processEvents()
            analyses=[]
            for oc in self.outcomes:
                gd=self._gd(oc)
                if gd: analyses.append(analyze_outcome(gd,list(gd.keys()),oc))
            prog.setValue(20)

            if analyses:
                add_h('Statistical Methods',2)
                for para in gen_methods_text(analyses).split('\n'):
                    if para.strip(): add_p(para.strip())
                add_p('')
                add_h('Statistical Results',2)
                for para in gen_results_text(analyses).split('\n'):
                    if para.strip(): add_p(para.strip())
            doc.add_page_break(); prog.setValue(30)

            # DESCRIPTIVE TABLE
            prog.setLabelText("Building tables..."); QApplication.processEvents()
            cl=['n','Mean +/- SD','Median','IQR','CV%']; rl=[]; cd=[]; sec=set(); ri=0
            for oc in self.outcomes:
                rl.append(oc); cd.append(['']*5); sec.add(ri); ri+=1
                for g in self.group_names:
                    vs=self.groups[g].get(oc,[])
                    if not vs: rl.append('  '+g); cd.append(['-']*5); ri+=1; continue
                    a=np.array(vs); m=np.mean(a); s=np.std(a,ddof=1); q1,q3=np.percentile(a,[25,75])
                    cv=s/m*100 if m else 0; rl.append('  '+g)
                    cd.append([str(len(a)),'{:.2f} +/- {:.2f}'.format(m,s),'{:.2f}'.format(np.median(a)),
                        '{:.2f} - {:.2f}'.format(q1,q3),'{:.1f}%'.format(cv)]); ri+=1
            add_h('Tables',2); add_tbl(cl,rl,cd,'Table 1. Descriptive Statistics',sec)
            prog.setValue(40)

            tbl_num=2
            for an in analyses:
                cl2=['Delta','p-adj','Sig','Hedges g [95% CI]','MDE']; rl2=[]; cd2=[]
                for r in an['pairs']:
                    rl2.append('{} vs {}'.format(r['g1'],r['g2']))
                    cd2.append(['{:+.2f}'.format(r['md']),'{:.4f}'.format(r['padj']),_sg(r['padj']),
                        '{:+.2f} [{:+.2f}, {:+.2f}]'.format(r['hg'],r['ci_lo'],r['ci_hi']),'{:.2f}'.format(r['mde'])])
                test='Mann-Whitney U' if an['nonp'] else 'Welch t-test'
                add_tbl(cl2,rl2,cd2,'Table {}. Pairwise: {} ({}, Bonferroni)'.format(tbl_num,an['oc'],test))
                tbl_num+=1
            prog.setValue(50)

            cd3=[]
            for g in self.group_names:
                row=[]
                for oc in self.outcomes:
                    vs=self.groups[g].get(oc,[])
                    if not vs: row.append('-')
                    else: a=np.array(vs); m=np.mean(a); row.append('{:.1f}%'.format(np.std(a,ddof=1)/m*100 if m else 0))
                cd3.append(row)
            add_tbl(self.outcomes,self.group_names,cd3,'Table {}. CV%'.format(tbl_num))
            doc.add_page_break(); prog.setValue(55)

            # FIGURES
            add_h('Figures',2); fig_num=1
            prog.setLabelText("Generating raincloud plots..."); QApplication.processEvents()
            for oi,oc in enumerate(self.outcomes):
                gd=self._gd(oc)
                if not gd: continue
                fig=gen_raincloud(gd,list(gd.keys()),self.colors,oc)
                add_fig(fig,'Figure {}. Raincloud plot: {}'.format(fig_num,oc)); fig_num+=1
                prog.setValue(55+int(oi*5)); QApplication.processEvents()

            prog.setLabelText("Generating bar plots..."); QApplication.processEvents()
            for oi,oc in enumerate(self.outcomes):
                gd=self._gd(oc)
                if not gd: continue
                fig=gen_barplot(gd,list(gd.keys()),self.colors,oc)
                add_fig(fig,'Figure {}. Bar plot: {}'.format(fig_num,oc)); fig_num+=1
                prog.setValue(75+int(oi*2)); QApplication.processEvents()

            prog.setLabelText("Generating scatter plots..."); QApplication.processEvents()
            for pi,(xk,yk) in enumerate(combinations(self.outcomes,2)):
                fig=gen_scatter(self.groups,self.group_names,self.colors,xk,yk)
                if fig: add_fig(fig,'Figure {}. {} vs {}'.format(fig_num,xk,yk)); fig_num+=1
                prog.setValue(85+int(pi*1.5)); QApplication.processEvents()

            prog.setLabelText("Generating heatmap..."); QApplication.processEvents()
            fig=gen_heatmap(self.groups,self.group_names,self.outcomes)
            if fig: add_fig(fig,'Figure {}. Correlation heatmap (pooled)'.format(fig_num))
            prog.setValue(95)

            prog.setLabelText("Saving..."); doc.save(p); prog.setValue(100)
            for f in tmp_files:
                try: os.unlink(f)
                except: pass
            QMessageBox.information(self,"Report Exported","Full report saved:\n{}\n\nTNR 12pt double-spaced (body)\nTNR 10pt single-spaced (tables)".format(p))
        except Exception as e:
            for f in tmp_files:
                try: os.unlink(f)
                except: pass
            QMessageBox.critical(self,"Error",str(e))

    # ════ INTERACTIVE PLOTS ════
    def _raincloud(self):
        if not self._ok(): return
        gd=self._gd(self.oc_box.currentText())
        if gd: self._show(gen_raincloud(gd,list(gd.keys()),self.colors,self.oc_box.currentText()))
    def _barplot(self):
        if not self._ok(): return
        gd=self._gd(self.oc_box.currentText())
        if gd: self._show(gen_barplot(gd,list(gd.keys()),self.colors,self.oc_box.currentText()))
    def _scatter_dlg(self):
        if not self._ok(): return
        dlg=QDialog(self); dlg.setWindowTitle("Scatter"); dlg.setFixedSize(350,200); lay=QVBoxLayout(dlg)
        lay.addWidget(QLabel("X:")); xb=QComboBox(); xb.addItems(self.outcomes); lay.addWidget(xb)
        lay.addWidget(QLabel("Y:")); yb=QComboBox(); yb.addItems(self.outcomes)
        if len(self.outcomes)>=2: yb.setCurrentIndex(1)
        lay.addWidget(yb); bb=QDialogButtonBox(QDialogButtonBox.Ok|QDialogButtonBox.Cancel)
        bb.accepted.connect(dlg.accept); bb.rejected.connect(dlg.reject); lay.addWidget(bb)
        if dlg.exec_()==QDialog.Accepted:
            fig=gen_scatter(self.groups,self.group_names,self.colors,xb.currentText(),yb.currentText())
            if fig: self._show(fig)
    def _heatmap(self):
        if not self._ok(): return
        fig=gen_heatmap(self.groups,self.group_names,self.outcomes)
        if fig: self._show(fig)

    # ════ INTERACTIVE TABLES ════
    def _tbl_desc(self):
        if not self._ok(): return
        cl=['n','Mean +/- SD','Median','IQR','CV%']; rl=[]; cd=[]; hl={}; sec=set(); ri=0
        for oc in self.outcomes:
            rl.append(oc); cd.append(['']*5); sec.add(ri); ri+=1
            for g in self.group_names:
                vs=self.groups[g].get(oc,[])
                if not vs: rl.append('  '+g); cd.append(['-']*5); ri+=1; continue
                a=np.array(vs); m=np.mean(a); s=np.std(a,ddof=1); q1,q3=np.percentile(a,[25,75])
                cv=s/m*100 if m else 0; rl.append('  '+g)
                cd.append([str(len(a)),'{:.2f} +/- {:.2f}'.format(m,s),'{:.2f}'.format(np.median(a)),
                    '{:.2f} - {:.2f}'.format(q1,q3),'{:.1f}%'.format(cv)])
                if cv>50: hl[(ri,4)]=True
                ri+=1
        self._show(gen_table_fig(cl,rl,cd,'Table 1. Descriptive Statistics',hl,sec))
    def _tbl_pw(self):
        if not self._ok(): return
        oc=self.oc_box.currentText(); gd=self._gd(oc)
        if not gd: return
        names=list(gd.keys()); pairs=list(combinations(range(len(names)),2)); ncp=len(pairs)
        cl=['Delta','p-adj','Sig','Hedges g [95% CI]','MDE']; rl=[]; cd=[]; hl={}
        for pi,(i,j) in enumerate(pairs):
            g1,g2=names[i],names[j]; a1,a2=gd[g1],gd[g2]; md=np.mean(a1)-np.mean(a2)
            U,pr=mannwhitneyu(a1,a2,alternative='two-sided'); pa=min(pr*ncp,1.)
            gv=hedges_g(a1,a2); ci=bootstrap_ci_g(a1,a2)
            sp_=np.sqrt((np.var(a1,ddof=1)+np.var(a2,ddof=1))/2); mv=mde_two(len(a1),len(a2),sp_)
            rl.append('{} vs {}'.format(g1,g2))
            cd.append(['{:+.2f}'.format(md),'{:.4f}'.format(pa),_sg(pa),'{:+.2f} [{:+.2f}, {:+.2f}]'.format(gv,ci[0],ci[1]),'{:.2f}'.format(mv)])
            if pa<.05:
                for c in range(5): hl[(pi,c)]=True
        self._show(gen_table_fig(cl,rl,cd,'Table 2. Pairwise: {} (Bonferroni)'.format(oc),hl))
    def _tbl_corr(self):
        if not self._ok(): return
        good=[oc for oc in self.outcomes if sum(len(self.groups[g].get(oc,[])) for g in self.group_names)>=5]
        if len(good)<2: return QMessageBox.warning(self,"","Need 2+ outcomes.")
        rd=[]
        for g in self.group_names:
            mx=max((len(self.groups[g].get(oc,[])) for oc in good),default=0)
            for i in range(mx):
                row={}; ok=True
                for oc in good:
                    vs=self.groups[g].get(oc,[])
                    if i<len(vs): row[oc]=vs[i]
                    else: ok=False; break
                if ok: rd.append(row)
        df=pd.DataFrame(rd); cd=[]; hl={}
        for i,o1 in enumerate(good):
            row=[]
            for j,o2 in enumerate(good):
                if i==j: row.append('--')
                else:
                    r,p=pearsonr(df[o1],df[o2]); sg='***' if p<.001 else '**' if p<.01 else '*' if p<.05 else ''
                    row.append('{:.3f}{}'.format(r,sg))
                    if p<.05: hl[(i,j)]=True
            cd.append(row)
        self._show(gen_table_fig(good,good,cd,'Table 3. Pearson Correlations (N={})'.format(len(df)),hl))
    def _tbl_cv(self):
        if not self._ok(): return
        cd=[]; hl={}
        for i,g in enumerate(self.group_names):
            row=[]
            for j,oc in enumerate(self.outcomes):
                vs=self.groups[g].get(oc,[])
                if not vs: row.append('-')
                else: a=np.array(vs); m=np.mean(a); cv=np.std(a,ddof=1)/m*100 if m else 0; row.append('{:.1f}%'.format(cv))
                if vs and cv>50: hl[(i,j)]=True
            cd.append(row)
        self._show(gen_table_fig(self.outcomes,self.group_names,cd,'Table 4. CV%',hl))

    # ════ SAVE/EXPORT ════
    def _save_fig(self):
        if not self.fig: return QMessageBox.warning(self,"","Generate a figure first.")
        p,_=QFileDialog.getSaveFileName(self,"Save","","PNG (*.png);;PDF (*.pdf);;TIFF (*.tiff);;SVG (*.svg)")
        if p: self.fig.savefig(p,dpi=300,bbox_inches='tight',facecolor='white',pad_inches=.1)
    def _export(self):
        if not self._ok(): return
        p,_=QFileDialog.getSaveFileName(self,"Export","","CSV (*.csv)")
        if not p: return
        rows=[]
        for oc in self.outcomes:
            for g in self.group_names:
                vs=self.groups[g].get(oc,[])
                if vs:
                    a=np.array(vs); m=np.mean(a)
                    rows.append({'Outcome':oc,'Group':g,'n':len(a),'Mean':round(m,4),'SD':round(np.std(a,ddof=1),4),
                        'Median':round(np.median(a),4),'CV_pct':round(np.std(a,ddof=1)/m*100,1) if m else 0})
        pd.DataFrame(rows).to_csv(p,index=False)

    def _welcome(self):
        self.txt.setPlainText(
            "  ======================================================\n"
            "       ProMetrix Analytics  v2.2\n"
            "       Courtesy of Dr. M S Omar  BDS  MSc\n"
            "       For Doers, Researchers and Innovators\n"
            "  ======================================================\n\n"
            "  1. Load data or Load Demo Data\n"
            "  2. Pick outcome / colors\n"
            "  3. Analyze:\n"
            "     - Run Full Statistics\n"
            "     - Plots: Raincloud, Bar, Scatter, Heatmap\n"
            "     - Tables: Descriptive, Pairwise, Correlation, CV%\n"
            "     - Generate Manuscript Text\n"
            "  4. EXPORT FULL REPORT:\n"
            "     One Word doc with EVERYTHING\n"
            "     Body: TNR 12pt double-spaced\n"
            "     Tables: TNR 10pt single-spaced\n\n"
            "  ------------------------------------------------\n")
    def _summary(self):
        self.txt.clear(); self._w("  DATA LOADED\n\n  Groups ({}):\n".format(len(self.group_names)))
        for g in self.group_names: self._w("    - {}\n".format(g))
        self._w("\n  Outcomes ({}):\n".format(len(self.outcomes)))
        for oc in self.outcomes:
            ns=[str(len(self.groups[g].get(oc,[]))) for g in self.group_names]
            self._w("    - {}  (n = {})\n".format(oc,', '.join(ns)))
        self._w("\n  -> Run Full Statistics or EXPORT FULL REPORT.\n")


_pm_window = None
def main():
    global _pm_window
    app=QApplication.instance()
    standalone=app is None
    if standalone: app=QApplication(sys.argv)
    _pm_window=ProMetrix()
    _pm_window.show(); _pm_window.raise_(); _pm_window.activateWindow()
    print("[ProMetrix Analytics] Launched.")
    if standalone: sys.exit(app.exec_())

if __name__=='__main__':
    main()
